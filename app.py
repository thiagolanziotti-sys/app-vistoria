import streamlit as st
import os
import io
from streamlit_js_eval import get_geolocation
from geopy.geocoders import Nominatim
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- FUNÇÕES XML (Cabeçalho/Rodapé do Word) ---
def create_element(name): return OxmlElement(name)
def create_attribute(element, name, value): element.set(qn(name), value)
def add_page_number(run):
    fldChar1 = create_element('w:fldChar'); create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText'); create_attribute(instrText, 'xml:space', 'preserve'); instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar'); create_attribute(fldChar2, 'w:fldCharType', 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)
def set_paragraph_border(paragraph, side="bottom"):
    p = paragraph._p; pPr = p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    border = OxmlElement(f'w:{side}')
    border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '8'); border.set(qn('w:space'), '1'); border.set(qn('w:color'), 'auto')
    pbdr.append(border); pPr.append(pbdr)

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="App Vistoria Imobiliária", layout="wide")

# --- LOGIN ---
def check_password():
    def password_entered():
        if st.session_state["username"] in st.secrets["passwords"] and \
           st.session_state["password"] == st.secrets["passwords"][st.session_state["username"]]:
            st.session_state["password_correct"] = True
            st.session_state["usuario_logado"] = st.session_state["username"] 
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.header("🔒 Acesso Restrito")
        st.text_input("Usuário", key="username")
        st.text_input("Senha", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.header("🔒 Acesso Restrito")
        st.text_input("Usuário", key="username")
        st.text_input("Senha", type="password", on_change=password_entered, key="password")
        st.error("Usuário ou senha incorretos.")
        return False
    else:
        return True

# --- GPS ---
def obter_endereco_por_coords(lat, lon):
    try:
        geolocator = Nominatim(user_agent="app_vistoria_thiago_v1")
        location = geolocator.reverse(f"{lat}, {lon}", addressdetails=True)
        return location.raw['address']
    except Exception as e:
        st.error(f"Erro GPS: {e}")
        return None

# --- APP PRINCIPAL ---
if check_password():
    
    if 'cliente' not in st.session_state: st.session_state.cliente = {}
    if 'avaliador' not in st.session_state: st.session_state.avaliador = {}
    if 'vistoria' not in st.session_state: st.session_state.vistoria = [] 
    
    usuario_atual = st.session_state.get("usuario_logado", "Avaliador")

    st.sidebar.success(f"Logado: {usuario_atual}")
    st.sidebar.title("Menu")
    
    # Lista das páginas para navegação
    PAGINAS = ["1. Avaliador", "2. Cliente e Imóvel", "3. Realizar Vistoria", "4. Relatórios"]
    
    # Se não tiver página definida, começa na primeira
    if "pagina_atual" not in st.session_state:
        st.session_state.pagina_atual = PAGINAS[0]

    # --- FUNÇÕES DE NAVEGAÇÃO (CORRIGIDAS: Sem st.rerun) ---
    def proxima_pagina():
        idx = PAGINAS.index(st.session_state.pagina_atual)
        if idx < len(PAGINAS) - 1:
            st.session_state.pagina_atual = PAGINAS[idx + 1]

    def pagina_anterior():
        idx = PAGINAS.index(st.session_state.pagina_atual)
        if idx > 0:
            st.session_state.pagina_atual = PAGINAS[idx - 1]

    # Menu Lateral (conectado à variável 'pagina_atual')
    escolha = st.sidebar.radio("Ir para:", PAGINAS, key="pagina_atual")

    # --- TELA 1 ---
    def tela_cadastro_avaliador():
        st.header("👷 Avaliador e Empresa")
        col1, col2 = st.columns(2)
        with st.form("form_avaliador"):
            with col1:
                st.subheader("Dados")
                nome = st.text_input("Nome", value=st.session_state.avaliador.get("nome", ""))
                titulo = st.selectbox("Título", ["Engenheiro Civil", "Arquiteto", "Corretor", "Perito"], index=0)
                registro = st.text_input("Registro (CREA/CAU)", value=st.session_state.avaliador.get("registro", ""))
            with col2:
                st.subheader("Visual")
                logo = st.file_uploader("Logo (Cabeçalho)", type=["png", "jpg"])
                if logo: st.image(logo, width=100)
                ass = st.file_uploader("Assinatura", type=["png", "jpg"])
            
            st.markdown("---")
            if st.form_submit_button("Salvar Configurações"):
                logo_final = logo if logo else st.session_state.avaliador.get("logo")
                ass_final = ass if ass else st.session_state.avaliador.get("assinatura")
                st.session_state.avaliador.update({"nome": nome, "titulo": titulo, "registro": registro, "assinatura": ass_final, "logo": logo_final})
                st.success("Salvo!")
            
        # --- BOTÕES DE NAVEGAÇÃO ---
        st.write("") 
        col_nav1, col_nav2 = st.columns([4, 1])
        with col_nav2:
            # CORREÇÃO: on_click
            st.button("Próximo ➡️", key="btn_nav_1", on_click=proxima_pagina)


    # --- TELA 2 ---
    def tela_cadastro_cliente_imovel():
        st.header("🏠 Cliente e Imóvel")
        col_c1, col_c2 = st.columns(2)
        with col_c1: st.session_state.cliente['nome'] = st.text_input("Cliente", st.session_state.cliente.get("nome", ""))
        with col_c2: st.session_state.cliente['contato'] = st.text_input("Contato", st.session_state.cliente.get("contato", ""))
        
        st.markdown("---")
        c_gps, c_btn = st.columns([3,1])
        c_gps.info("Ative a localização do navegador.")
        with c_btn:
            loc = get_geolocation()
            if loc and st.button("📍 GPS"):
                end = obter_endereco_por_coords(loc['coords']['latitude'], loc['coords']['longitude'])
                if end:
                    st.session_state.cliente.update({'rua': end.get('road',''), 'numero': end.get('house_number',''), 'bairro': end.get('suburb', end.get('neighbourhood','')), 'cidade': end.get('city', end.get('town','')), 'estado': end.get('state','')})
                    st.rerun()

        c1, c2, c3 = st.columns([3, 1, 1])
        st.session_state.cliente['rua'] = c1.text_input("Rua", st.session_state.cliente.get("rua", ""))
        st.session_state.cliente['numero'] = c2.text_input("Nº", st.session_state.cliente.get("numero", ""))
        st.session_state.cliente['complemento'] = c3.text_input("Compl.", st.session_state.cliente.get("complemento", ""))
        
        c4, c5, c6 = st.columns([2, 2, 1])
        st.session_state.cliente['bairro'] = c4.text_input("Bairro", st.session_state.cliente.get("bairro", ""))
        st.session_state.cliente['cidade'] = c5.text_input("Cidade", st.session_state.cliente.get("cidade", ""))
        st.session_state.cliente['estado'] = c6.text_input("UF", st.session_state.cliente.get("estado", ""))

        st.markdown("---")
        t1, t2, t3 = st.columns(3)
        st.session_state.cliente['tipo_imovel'] = t1.selectbox("Tipo", ["Casa", "Apto", "Comercial", "Galpão", "Terreno"], index=0)
        st.session_state.cliente['matricula'] = t2.text_input("Matrícula", st.session_state.cliente.get("matricula", ""))
        st.session_state.cliente['estado_global'] = t3.selectbox("Estado Global", ["Novo", "Bom", "Regular", "Ruim", "Crítico"], index=1)

        a1, a2, a3 = st.columns(3)
        st.session_state.cliente['area_const'] = a1.number_input("Área Const (m²)", value=float(st.session_state.cliente.get("area_const", 0)))
        st.session_state.cliente['area_terreno'] = a2.number_input("Área Terreno (m²)", value=float(st.session_state.cliente.get("area_terreno", 0)))
        st.session_state.cliente['quartos'] = a3.number_input("Quartos", value=int(st.session_state.cliente.get("quartos", 0)))

        # --- BOTÕES DE NAVEGAÇÃO E SALVAR ---
        st.markdown("---")
        c_voltar, c_salvar, c_prox = st.columns([1, 2, 1])
        
        with c_voltar:
            st.button("⬅️ Voltar", key="btn_voltar_2", on_click=pagina_anterior)
        
        with c_salvar:
            if st.button("💾 SALVAR DADOS", type="primary", use_container_width=True):
                st.success("✅ Salvo!")
        
        with c_prox:
            st.button("Próximo ➡️", key="btn_prox_2", on_click=proxima_pagina)


    # --- TELA 3 ---
    def tela_vistoria_fotos():
        st.header("📸 Vistoria")
        with st.container():
            c1, c2 = st.columns([2, 1])
            amb = c1.text_input("Ambiente")
            cond = c2.selectbox("Condição", ["Ótimo", "Bom", "Regular", "Ruim", "Péssimo"])
            d1, d2 = st.columns(2)
            desc = d1.text_area("1. Descrição")
            analise = d2.text_area("2. Patologia/Análise")
            
            t_cam, t_arq = st.tabs(["Câmera", "Arquivo"])
            foto = None
            with t_cam: 
                f = st.camera_input("Foto")
                if f: foto = f
            with t_arq: 
                f = st.file_uploader("Upload", type=["jpg","png"])
                if f: foto = f

            if st.button("➕ Adicionar", type="primary"):
                if amb and foto:
                    st.session_state.vistoria.append({"ambiente": amb, "condicao": cond, "descricao": desc, "analise": analise, "foto": foto.getvalue()})
                    st.success("Adicionado!")
                else: st.error("Nome e Foto obrigatórios.")

        st.divider()
        for i, item in enumerate(st.session_state.vistoria):
            with st.expander(f"{i+1}. {item['ambiente']} ({item['condicao']})"):
                ci, ct = st.columns([1,2])
                ci.image(item['foto'], use_container_width=True)
                ct.write(f"**Desc:** {item['descricao']}\n\n**Análise:** {item['analise']}")
                if ct.button("Excluir", key=f"del_{i}"):
                    st.session_state.vistoria.pop(i); st.rerun()

        # --- BOTÕES DE NAVEGAÇÃO ---
        st.write("")
        st.markdown("---")
        c_voltar, c_prox = st.columns([1, 1])
        
        with c_voltar:
            st.button("⬅️ Voltar", key="btn_voltar_3", on_click=pagina_anterior)
                
        with c_prox:
            # Usamos on_click para avançar, mas colocamos validação visual antes
            st.button("✅ Concluir e Gerar Relatório ➡️", type="primary", key="btn_prox_3", on_click=proxima_pagina)


    # --- TELA 4 ---
    def tela_relatorios():
        st.header("📄 Relatório Final")
        if not st.session_state.vistoria: 
            st.warning("⚠️ Você ainda não realizou vistorias. Volte e adicione fotos.")
            if st.button("⬅️ Voltar para Vistoria"):
                pagina_anterior()
            return

        def gerar_docx_profissional():
            doc = Document()
            sec = doc.sections[0]
            sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
            sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
            sec.header_distance = Inches(0.3); sec.footer_distance = Inches(0.3)

            # Cabeçalho
            header = sec.header
            ht = header.add_table(1, 2, width=Inches(6)); ht.autofit = False
            ht.columns[0].width = Inches(1.5); ht.columns[1].width = Inches(5.0)
            if st.session_state.avaliador.get("logo"):
                try: ht.cell(0,0).paragraphs[0].add_run().add_picture(io.BytesIO(st.session_state.avaliador["logo"].getvalue()), height=Inches(0.6))
                except: pass
            hp = ht.cell(0,1).paragraphs[0]; hp.alignment = 2
            hp.add_run("VISTORIA TÉCNICA\n").bold = True
            hp.add_run("Página "); add_page_number(hp.add_run())
            set_paragraph_border(header.add_paragraph(), "bottom")

            # Rodapé
            footer = sec.footer
            set_paragraph_border(footer.paragraphs[0], "top")
            fp = footer.add_paragraph(); fp.alignment = 1
            fp.add_run(f"{st.session_state.avaliador.get('nome', '')} | {st.session_state.avaliador.get('registro', '')}")

            # Conteúdo
            doc.add_heading('RELATÓRIO TÉCNICO', 1)
            cli = st.session_state.cliente
            t = doc.add_table(rows=5, cols=2); t.style = 'Table Grid'
            def sr(i,l,v): t.cell(i,0).text=l; t.cell(i,0).paragraphs[0].runs[0].bold=True; t.cell(i,1).text=str(v)
            sr(0,"Cliente:", cli.get('nome',''))
            sr(1,"Endereço:", f"{cli.get('rua','')}, {cli.get('numero','')} {cli.get('complemento','')}\n{cli.get('bairro','')} - {cli.get('cidade','')}/{cli.get('estado','')}")
            sr(2,"Tipo:", cli.get('tipo_imovel',''))
            sr(3,"Matrícula:", cli.get('matricula',''))
            sr(4,"Dados:", f"Área: {cli.get('area_const',0)}m² | Quartos: {cli.get('quartos',0)}")
            
            doc.add_paragraph(" ")
            doc.add_heading('VISTORIA FOTOGRÁFICA', 1)
            
            for i, item in enumerate(st.session_state.vistoria):
                ti = doc.add_table(rows=3, cols=1); ti.style = 'Table Grid'; ti.autofit = False
                pp = ti.cell(0,0).paragraphs[0]; pp.alignment = 1
                try: 
                    pp.add_run().add_picture(io.BytesIO(item['foto']), width=Inches(3.2))
                except: pp.text = "[Erro Imagem]"
                ti.cell(1,0).text = f"Item {i+1}: {item['ambiente']} | {item['condicao'].upper()}"
                ti.cell(2,0).text = f"Desc: {item['descricao']}\nAnálise: {item['analise']}"
                doc.add_paragraph(" ")

            # Assinatura
            doc.add_paragraph(" "); doc.add_paragraph(" ")
            pa = doc.add_paragraph(); pa.alignment = 1
            if st.session_state.avaliador.get("assinatura"):
                try: pa.add_run().add_picture(io.BytesIO(st.session_state.avaliador["assinatura"].getvalue()), width=Inches(2.0))
                except: pass
            pa.add_run(f"\n_______________________\n{st.session_state.avaliador.get('nome', '')}")

            buffer = io.BytesIO(); doc.save(buffer); buffer.seek(0)
            return buffer

        # BOTÕES
        st.markdown("---")
        
        # Geramos o buffer ANTES do botão para ele estar pronto para o download
        buf = gerar_docx_profissional()
        nome_arq = f"Vistoria - {st.session_state.cliente.get('nome', 'Cliente')}.docx"
            
        col_d1, col_d2 = st.columns([3, 1])
        with col_d1:
            st.download_button(
                label="📥 BAIXAR RELATÓRIO DOCX",
                data=buf,
                file_name=nome_arq,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
        
        with col_d2:
            st.button("⬅️ Voltar", key="btn_voltar_4", on_click=pagina_anterior)

    # Lógica de exibição das telas baseada no session_state
    if escolha == "1. Avaliador": tela_cadastro_avaliador()
    elif escolha == "2. Cliente e Imóvel": tela_cadastro_cliente_imovel()
    elif escolha == "3. Realizar Vistoria": tela_vistoria_fotos()
    elif escolha == "4. Relatórios": tela_relatorios()